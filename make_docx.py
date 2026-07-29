"""Build a styled .docx from cv.html — same content, Word-native formatting.

Usage: python make_docx.py [output.docx]
"""
import re
import sys
from html.parser import HTMLParser

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SERIF, SANS = "Charter", "Helvetica Neue"
INK      = RGBColor(0x14, 0x17, 0x1A)
INK_SOFT = RGBColor(0x3A, 0x40, 0x46)
INK_MUTE = RGBColor(0x71, 0x7A, 0x85)
ACCENT   = RGBColor(0x7A, 0x1F, 0x2B)
BODY_PT  = 9.2

# ── html → flat run list ────────────────────────────────────────────────
class Runs(HTMLParser):
    """Collect (text, bold, italic, underline) tuples from an HTML fragment."""
    def __init__(self):
        super().__init__()
        self.out, self.b, self.i, self.u = [], 0, 0, 0
    def handle_starttag(self, tag, attrs):
        cls = dict(attrs).get("class", "")
        if tag in ("b", "strong") or "lead" in cls: self.b += 1
        if tag in ("i", "em") or "venue" in cls:    self.i += 1
        if "me" in cls.split():                     self.u += 1
        if "yr" in cls.split():                     self.i += 1
    def handle_endtag(self, tag):
        pass  # depth tracked loosely below via re-parse of balanced fragments
    def handle_data(self, data):
        self.out.append((data, self.b > 0, self.i > 0, self.u > 0))

TAG = re.compile(r"<(/?)(\w+)([^>]*)>")

def runs_of(html):
    """Parse a fragment into styled runs, tracking nesting properly."""
    out, pos = [], 0
    bold = ital = und = 0
    for m in TAG.finditer(html):
        text = html[pos:m.start()]
        if text:
            out.append((unesc(text), bold > 0, ital > 0, und > 0))
        closing, tag, attrs = m.group(1), m.group(2).lower(), m.group(3)
        cls = re.search(r'class="([^"]*)"', attrs)
        cls = cls.group(1).split() if cls else []
        d = -1 if closing else 1
        if tag in ("b", "strong") or "lead" in cls:      bold += d
        if tag in ("i", "em") or "venue" in cls or "yr" in cls: ital += d
        if "me" in cls:                                  und += d
        pos = m.end()
    if html[pos:]:
        out.append((unesc(html[pos:]), bold > 0, ital > 0, und > 0))
    return [(t, b, i, u) for t, b, i, u in out if t]

def unesc(t):
    for a, b in (("&amp;", "&"), ("&nbsp;", " "), ("&lt;", "<"), ("&gt;", ">"),
                 ("&times;", "×"), ("&rho;", "ρ")):
        t = t.replace(a, b)
    return re.sub(r"\s+", " ", t)

# ── docx helpers ────────────────────────────────────────────────────────
def style_run(r, size=BODY_PT, font=SERIF, color=INK, bold=False, italic=False,
              underline=False, caps=False, spacing=None):
    r.font.name, r.font.size, r.font.color.rgb = font, Pt(size), color
    r.bold, r.italic, r.underline = bold, italic, underline
    if caps:
        r.font.all_caps = True
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), font)
    if spacing:
        sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(int(spacing * 20))); rpr.append(sp)
    return r

def para(doc, before=0, after=0, line=1.36, indent=0, hang=None, tabs=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    pf.line_spacing = line
    if indent:
        pf.left_indent = Inches(indent)
    if hang:
        pf.first_line_indent = Inches(-hang)
    if tabs:
        for pos, align in tabs:
            p.paragraph_format.tab_stops.add_tab_stop(Inches(pos), align)
    return p

def bottom_border(p, color="CCD2D8", sz=6):
    pPr = p._element.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), color)
    bdr.append(b); pPr.append(bdr)

def section_head(doc, text):
    p = para(doc, before=10, after=2.5)
    style_run(p.add_run(text.upper()), size=7.7, font=SANS, color=INK_SOFT,
              bold=True, spacing=1.3)
    bottom_border(p)
    return p

def emit_runs(p, html, size=BODY_PT):
    for t, b, i, u in runs_of(html):
        style_run(p.add_run(t), size=size, color=INK if b else INK_SOFT,
                  bold=b, italic=i, underline=u)

def bullet(doc, html, indent=0.17, size=BODY_PT):
    p = para(doc, before=3.4, after=0, indent=indent, hang=0.17)
    style_run(p.add_run("–\t"), size=size, color=RGBColor(0xA9, 0xB1, 0xBA))
    p.paragraph_format.tab_stops.add_tab_stop(Inches(indent), WD_TAB_ALIGNMENT.LEFT)
    emit_runs(p, html, size=size)
    return p

# ── read source ─────────────────────────────────────────────────────────
src = open("cv.html", encoding="utf-8").read()
BODY = src[src.index("<body>"):]
RIGHT = 7.3  # right tab position (inches) for dates

doc = Document()
sec = doc.sections[0]
sec.top_margin, sec.bottom_margin = Inches(0.45), Inches(0.4)
sec.left_margin, sec.right_margin = Inches(0.6), Inches(0.6)
usable = 8.5 - 1.2
RIGHT = usable

# header ────────────────────────────────────────────────────────────────
p = para(doc, after=0, line=1.0)
style_run(p.add_run("Xinyu (Brian) Guo"), size=22.5, font=SERIF, color=INK, bold=True)

p = para(doc, before=4, after=0, line=1.0)
style_run(p.add_run("GENOMIC FOUNDATION MODELS  ·  CANCER GENOMICS  ·  SCIENTIFIC AI AGENTS"),
          size=8, font=SANS, color=ACCENT, bold=True, spacing=.65)

p = para(doc, before=3.5, after=0, line=1.0)
style_run(p.add_run("xinyug@usc.edu   ·   (314) 680-8961   ·   Open to relocation   ·   "
                    "xinyuguo.com   ·   github.com/Thewhey-Brian"),
          size=8.75, color=INK_SOFT)
bottom_border(p, color="14171A", sz=10)

# education ─────────────────────────────────────────────────────────────
section_head(doc, "Education")
for degree, rest, when in [
    ("Ph.D., Computational Biology & Bioinformatics",
     "  University of Southern California · 3.95/4.0 · Viterbi Fellow", "Aug 2022 – 2026"),
    ("M.S., Biostatistics",
     "  Johns Hopkins University · 3.97/4.0 · Delta Omega Honor Society", "Aug 2020 – May 2022"),
    ("B.A., Mathematics & Computer Science",
     "  Washington University in St. Louis · 3.89/4.0 · Cum Laude", "Aug 2018 – May 2020"),
]:
    p = para(doc, before=3.4, after=0, tabs=[(RIGHT, WD_TAB_ALIGNMENT.RIGHT)])
    style_run(p.add_run(degree), size=9.4, font=SANS, color=INK, bold=True)
    head, _, tail = rest.rpartition("· ")
    style_run(p.add_run(head + "· "), size=8.9, color=INK_SOFT)
    style_run(p.add_run(tail), size=8.9, color=INK_SOFT, italic=True)
    p.add_run("\t")
    style_run(p.add_run(when), size=8.2, font=SANS, color=INK_MUTE)

# experience / projects / publications / skills ─────────────────────────
ENTRY = re.compile(
    r'<div class="erow">\s*<div class="role">(.*?)</div>\s*'
    r'<div class="when">(.*?)</div>\s*</div>\s*'
    r'<div class="sub">(.*?)</div>\s*<ul>(.*?)</ul>', re.S)
LI = re.compile(r"<li>(.*?)</li>", re.S)

section_head(doc, "Experience")
for role, when, sub, items in ENTRY.findall(BODY):
    p = para(doc, before=7, after=0, tabs=[(RIGHT, WD_TAB_ALIGNMENT.RIGHT)])
    for t, b, i, u in runs_of(role):
        style_run(p.add_run(t), size=9.7, font=SANS,
                  color=INK if "—" not in t else INK_SOFT, bold=True)
    p.add_run("\t")
    style_run(p.add_run(unesc(when)), size=8.2, font=SANS, color=INK_MUTE)
    p = para(doc, before=0.5, after=0)
    style_run(p.add_run(unesc(re.sub("<[^>]+>", "", sub))), size=8.8,
              color=INK_MUTE, italic=True)
    for li in LI.findall(items):
        bullet(doc, li)

def loose_section(title, html_block):
    section_head(doc, title)
    for li in LI.findall(html_block):
        bullet(doc, li)

proj = BODY[BODY.index("<h2>Selected Projects</h2>"):BODY.index("<h2>Selected Publications</h2>")]
loose_section("Selected Projects", proj)

# publications
section_head(doc, "Selected Publications")
pubs = BODY[BODY.index('<ol class="pubs">'):BODY.index("</ol>")]
for n, li in enumerate(LI.findall(pubs), 1):
    p = para(doc, before=2.6, after=0, indent=0.2, hang=0.2)
    style_run(p.add_run(f"{n}.\t"), size=7.6, font=SANS, color=INK_MUTE)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(0.2), WD_TAB_ALIGNMENT.LEFT)
    emit_runs(p, li, size=8.95)

# skills
section_head(doc, "Technical Expertise")
KV = re.compile(r'<span class="k">(.*?)</span>\s*&nbsp;<span class="v">(.*?)</span>', re.S)
skills = BODY[BODY.index("<h2>Technical Expertise</h2>"):]
for k, v in KV.findall(skills):
    p = para(doc, before=3.4, after=0)
    style_run(p.add_run(unesc(k).upper() + "  "), size=7.8, font=SANS,
              color=ACCENT, bold=True, spacing=.4)
    style_run(p.add_run(unesc(re.sub("<[^>]+>", "", v))), size=BODY_PT, color=INK_SOFT)

out = sys.argv[1] if len(sys.argv) > 1 else "dist/Xinyu_Guo_CV.docx"
doc.save(out)
print("wrote", out)
